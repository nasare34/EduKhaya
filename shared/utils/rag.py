import os
import uuid
import hashlib
import time
import chromadb
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv
from typing import Generator

load_dotenv()

CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "./chroma_db")

_embedding_model = None

def get_embedding_model():
    global _embedding_model
    if _embedding_model is None:
        _embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    return _embedding_model

def get_chroma_client():
    return chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)

def get_collection_name(user_id: int, subject: str) -> str:
    raw = f"user_{user_id}_{subject.lower().replace(' ', '_')}"
    return f"col_{hashlib.md5(raw.encode()).hexdigest()[:16]}"

def make_chunk_id(user_id: int, filename: str, chunk_index: int, chunk_text: str) -> str:
    """
    Generate a guaranteed-unique chunk ID.
    Combines user_id + filename + chunk index + first 40 chars of text + uuid fragment.
    This prevents collisions when:
      - Same document is re-uploaded
      - Different documents share identical paragraphs
      - Same user uploads multiple files to the same subject
    """
    base = f"{user_id}_{filename}_{chunk_index}_{chunk_text[:40]}"
    stable_hash = hashlib.md5(base.encode()).hexdigest()[:12]
    unique_suffix = uuid.uuid4().hex[:8]
    return f"{stable_hash}_{unique_suffix}"


# ─── Document loaders ─────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF. Falls back to OCR if text layer is empty."""
    try:
        from langchain_community.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        text = "\n\n".join([d.page_content for d in docs if d.page_content.strip()])
        if len(text.strip()) > 100:
            return text
    except Exception:
        pass

    # Fallback: OCR the PDF pages
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import pytesseract
        import io

        doc = fitz.open(file_path)
        pages_text = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text = pytesseract.image_to_string(img)
            if ocr_text.strip():
                pages_text.append(ocr_text)
        doc.close()
        return "\n\n".join(pages_text)
    except Exception as e:
        raise ValueError(f"Could not extract text from PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word document — handles all content including tables."""
    try:
        import docx
        document = docx.Document(file_path)
        parts = []

        # Extract paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract table content
        for table in document.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n\n".join(parts)
    except Exception:
        # Fallback to langchain loader
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(file_path)
            docs = loader.load()
            return "\n\n".join([d.page_content for d in docs if d.page_content.strip()])
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {e}")


def extract_text_from_image(file_path: str) -> str:
    """OCR text from image files (PNG, JPG, JPEG, TIFF, BMP)."""
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)
        if not text.strip():
            raise ValueError("No text found in image after OCR.")
        return text
    except ImportError:
        raise ValueError("Pillow and pytesseract are required for image OCR. Run: pip install pillow pytesseract")
    except Exception as e:
        raise ValueError(f"Image OCR failed: {e}")


def extract_text_from_txt(file_path: str) -> str:
    """Extract plain text. Tries UTF-8 first, then latin-1."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def load_document(file_path: str) -> tuple:
    """
    Load any supported file and return (full_text, page_count).
    Supported: .pdf, .docx, .doc, .txt, .md, .png, .jpg, .jpeg, .tiff, .bmp
    Does NOT require documents to have any particular structure or IDs.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
        # Estimate page count from text length
        page_count = max(1, len(text) // 2000)
        return text, page_count

    elif ext in ('.docx', '.doc'):
        text = extract_text_from_docx(file_path)
        page_count = max(1, len(text) // 2000)
        return text, page_count

    elif ext in ('.txt', '.md'):
        text = extract_text_from_txt(file_path)
        return text, 1

    elif ext in ('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp'):
        text = extract_text_from_image(file_path)
        return text, 1

    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported formats: PDF, DOCX, DOC, TXT, MD, PNG, JPG, JPEG, TIFF, BMP"
        )


# ─── Silent ingest ────────────────────────────────────────────────────────────

def ingest_document(file_path: str, user_id: int, subject: str, grade_level: str = "") -> dict:
    """Load, chunk, embed and store. Returns summary dict."""
    filename = os.path.basename(file_path)
    full_text, page_count = load_document(file_path)

    if not full_text.strip():
        raise ValueError("The document appears to be empty or has no readable text.")

    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    chunks = splitter.split_text(full_text)

    if not chunks:
        raise ValueError("Could not split document into chunks. The file may be too short or unreadable.")

    model = get_embedding_model()
    embeddings = model.encode(chunks).tolist()

    client = get_chroma_client()
    collection_name = get_collection_name(user_id, subject)
    collection = client.get_or_create_collection(
        name=collection_name,
        metadata={"subject": subject, "grade_level": grade_level, "user_id": str(user_id)}
    )

    # Unique IDs — never collide even with identical content
    ids = [make_chunk_id(user_id, filename, i, chunk) for i, chunk in enumerate(chunks)]
    metadatas = [{"source": filename, "chunk_index": i, "subject": subject} for i in range(len(chunks))]

    collection.add(documents=chunks, embeddings=embeddings, ids=ids, metadatas=metadatas)

    return {
        "collection_name": collection_name,
        "chunk_count": len(chunks),
        "char_count": len(full_text),
        "page_count": page_count
    }


# ─── Streaming ingest with step-by-step progress ─────────────────────────────

def ingest_document_stream(
    file_path: str,
    user_id: int,
    subject: str,
    grade_level: str = ""
) -> Generator[dict, None, None]:
    """
    Same as ingest_document but yields progress dicts at every meaningful step.
    Each dict: { step, total_steps, label, detail, progress (0-100), done, error }
    """
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].upper().lstrip('.')

    try:
        # ── Step 1: Reading the file ──────────────────────────────────────────
        yield {
            "step": 1, "total_steps": 5,
            "label": "Reading your file",
            "detail": f"Opening {filename} and extracting all text content from the {ext} document. This works even if the document has no structure, headings, or IDs.",
            "progress": 5, "done": False, "error": None
        }

        full_text, page_count = load_document(file_path)

        if not full_text.strip():
            raise ValueError("The document appears to be empty or has no readable text.")

        char_count = len(full_text)

        yield {
            "step": 1, "total_steps": 5,
            "label": "File read successfully",
            "detail": f"Extracted {char_count:,} characters across {page_count} page(s). The AI can now see the full contents of your document.",
            "progress": 20, "done": False, "error": None
        }

        # ── Step 2: Chunking ──────────────────────────────────────────────────
        yield {
            "step": 2, "total_steps": 5,
            "label": "Splitting into knowledge chunks",
            "detail": "Breaking the text into overlapping 800-character segments. The 100-character overlap ensures no important sentence is cut in half between chunks. No document structure or IDs are required.",
            "progress": 25, "done": False, "error": None
        }

        splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
        chunks = splitter.split_text(full_text)
        chunk_count = len(chunks)

        if not chunks:
            raise ValueError("Could not split document into chunks. The file may be too short or unreadable.")

        yield {
            "step": 2, "total_steps": 5,
            "label": f"Created {chunk_count} knowledge chunks",
            "detail": f"Your {char_count:,}-character document became {chunk_count} searchable chunks. Each chunk gets its own unique ID generated from its position and content — so the same text uploaded twice will never cause an ID conflict.",
            "progress": 38, "done": False, "error": None
        }

        # ── Step 3: Embedding ─────────────────────────────────────────────────
        yield {
            "step": 3, "total_steps": 5,
            "label": "Loading AI embedding model",
            "detail": "Initialising the MiniLM-L6 sentence transformer. This converts text into 384-dimensional vectors capturing semantic meaning — not just keywords.",
            "progress": 42, "done": False, "error": None
        }

        model = get_embedding_model()

        yield {
            "step": 3, "total_steps": 5,
            "label": f"Embedding {chunk_count} chunks into vectors",
            "detail": f"Converting each chunk into a list of 384 numbers. Chunks with similar meaning will have vectors that are close together — enabling smart retrieval later.",
            "progress": 48, "done": False, "error": None
        }

        # Embed in batches and report progress
        batch_size = max(1, min(32, chunk_count // 4 + 1))
        all_embeddings = []
        for i in range(0, chunk_count, batch_size):
            batch = chunks[i:i + batch_size]
            batch_vecs = model.encode(batch).tolist()
            all_embeddings.extend(batch_vecs)
            pct_done = min(int((len(all_embeddings) / chunk_count) * 100), 100)
            embed_progress = 48 + int(pct_done * 0.22)
            yield {
                "step": 3, "total_steps": 5,
                "label": f"Embedding chunks ({len(all_embeddings)}/{chunk_count})",
                "detail": f"Vectorised {len(all_embeddings)} of {chunk_count} chunks.",
                "progress": embed_progress, "done": False, "error": None
            }

        yield {
            "step": 3, "total_steps": 5,
            "label": "All chunks vectorised",
            "detail": f"Successfully converted all {chunk_count} chunks into embedding vectors.",
            "progress": 70, "done": False, "error": None
        }

        # ── Step 4: Storing in ChromaDB ───────────────────────────────────────
        yield {
            "step": 4, "total_steps": 5,
            "label": "Storing in knowledge base (ChromaDB)",
            "detail": f"Saving {chunk_count} chunks into your {subject} collection with collision-safe unique IDs. Each ID combines: filename + chunk position + content hash + random suffix.",
            "progress": 75, "done": False, "error": None
        }

        client = get_chroma_client()
        collection_name = get_collection_name(user_id, subject)
        collection = client.get_or_create_collection(
            name=collection_name,
            metadata={"subject": subject, "grade_level": grade_level, "user_id": str(user_id)}
        )

        # Collision-safe unique IDs
        ids = [make_chunk_id(user_id, filename, i, chunk) for i, chunk in enumerate(chunks)]
        metadatas = [{"source": filename, "chunk_index": i, "subject": subject} for i in range(chunk_count)]

        collection.add(
            documents=chunks,
            embeddings=all_embeddings,
            ids=ids,
            metadatas=metadatas
        )

        yield {
            "step": 4, "total_steps": 5,
            "label": "Knowledge base updated",
            "detail": f"All {chunk_count} chunks are now indexed in your {subject} collection. You can safely upload the same document again — IDs will never clash.",
            "progress": 90, "done": False, "error": None
        }

        # ── Step 5: Done ──────────────────────────────────────────────────────
        yield {
            "step": 5, "total_steps": 5,
            "label": "Ready to use!",
            "detail": f'The AI has fully learned from "{filename}". When you generate content for {subject}, it will retrieve the most relevant chunks from this document.',
            "progress": 100,
            "done": True,
            "error": None,
            "collection_name": collection_name,
            "chunk_count": chunk_count,
            "char_count": char_count,
            "page_count": page_count
        }

    except Exception as e:
        yield {
            "step": -1, "total_steps": 5,
            "label": "Processing failed",
            "detail": str(e),
            "progress": 0, "done": True, "error": str(e)
        }


# ─── Retrieval ────────────────────────────────────────────────────────────────

def retrieve_context(query: str, user_id: int, subject: str, top_k: int = 5) -> str:
    try:
        client = get_chroma_client()
        collection_name = get_collection_name(user_id, subject)
        collection = client.get_collection(name=collection_name)
        model = get_embedding_model()
        query_embedding = model.encode([query]).tolist()
        results = collection.query(query_embeddings=query_embedding, n_results=top_k)
        documents = results.get('documents', [[]])[0]
        return "\n\n---\n\n".join(documents)
    except Exception:
        return ""

def list_user_collections(user_id: int) -> list:
    client = get_chroma_client()
    all_collections = client.list_collections()
    user_collections = []
    for col in all_collections:
        meta = col.metadata or {}
        if str(meta.get("user_id")) == str(user_id):
            user_collections.append({
                "name": col.name,
                "subject": meta.get("subject", ""),
                "grade_level": meta.get("grade_level", "")
            })
    return user_collections

def delete_collection(collection_name: str):
    client = get_chroma_client()
    client.delete_collection(name=collection_name)
